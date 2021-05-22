#This program was made to add changes to multiple docx files from just one python program. 

import docx
from docx.shared import Pt
from datetime import datetime
import calendar
import os

def ticket_to_Istanbul (myname, check_in_date, check_out_date, booking):
    os.getcwd()

    doc = docx.Document('source/ticket to Istanbul.docx')
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Verdana'
    font.size = Pt(7.5)

    #name
    all_para = doc.paragraphs
    all_para[0].text = myname
    run = all_para[0].runs
    fonts = run[0].font
    fonts.name = 'Verdana'
    fonts.size = Pt(11)
    fonts.bold = True

    #booking Ref
    booking_ref = all_para[1].runs
    booking_ref[3].text = f": {booking}"
    #Issue Date
    issue_date = all_para[1].runs
    issue_date[24].text = datetime.today().strftime('%d')
    issue_date[26].text = datetime.today().strftime('%m')
    issue_date[27].text = "/20" + datetime.today().strftime('%y')
    issue_date[28].text = ""

    #tarih date1
    tables = doc.tables

    rows = tables[0].rows
    cells = rows[1].cells
    cells[3].text = check_in_date

    row = tables[1].rows
    cell = row[1].cells
    cell[3].text = check_out_date

    doc.save('result/ticket-to-Istanbul.docx')

def PrivatPolice (address, name, police_nr, birthdate):
    os.getcwd()

    doc = docx.Document('source/__PrivatPolice.docx')
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(9)

    #first appearance of name
    all_para = doc.paragraphs
    name1 = all_para[8].runs
    name1[0].text = name
    name1[2].text = address
    i = 3
    while i<8:
        name1[i].text = ""
        i += 1

    #2nd appearance of name
    tables = doc.tables
    rows = tables[0].rows
    cells = rows[1].cells
    cells[0].text = name

    #police nr
    all_para[5].text = f"Policenr. {police_nr}"

    #birthdate
    cells[1].text = f"{birthdate}-####"
    doc.save('result/__PrivatPolice.docx')

def agan_Hotel(name, start_date, end_date, no_of_nights, start_month, end_month):
    os.getcwd()

    doc = docx.Document('source/agan_Hotel.docx')
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Segoe UI'
    font.size = Pt(7.5)

    #date from, date till and number of nights
    all_para = doc.paragraphs
    all_para[17].text = f"{start_date}          {end_date}             {no_of_nights}"
    run = all_para[17].runs
    fonts = run[0].font
    fonts.name = 'Arial'
    fonts.size = Pt(20)
    fonts.bold = True

    #start month
    start_month_run = all_para[21].runs
    start_month_run[0].text = calendar.month_name[int(start_month)]

    #end month
    end_month_run = all_para[24].runs
    end_month_run[1].text = calendar.month_name[int(end_month)]
    doc.save('result/agan_Hotel.docx')

check_in_date_wrong = True
check_out_date_wrong = True
not_done = True
police_nr_not_done = True
name = input("Enter the name: ")
address = input("Please Enter the Address: ")

while check_in_date_wrong:
    print("Please Enter Check-In Date: ")
    date1 = input("Enter day: ")
    month1 = input("Enter month: ")
    year1 = input("Enter year: ")
    if 1 <= int(date1) and int(date1) <= 31 and 1 <= int(month1) and int(month1) <= 12:
        check_in_date_wrong = False
    else:
        print("ERROR: Invalid Input. Try Again")

check_in_date = f"{date1}/{month1}/{year1}"

while check_out_date_wrong:
    print("\nPlease Enter Check-Out Date: ")
    date2 = input("Enter day: ")
    month2 = input("Enter month: ")
    year2 = input("Enter year: ")
    if 1 <= int(date2) and int(date2) <= 31 and 1 <= int(month2) and int(month2) <= 12 and year2 >= year1:

        check_out_date_wrong = False

    else:
        print("ERROR: Invalid Input. Try Again")

    if year2 == year1 and int(month1 > month2):
        print("Check out date cannot be before check in date")
        check_out_date_wrong = True
    elif year2 == year1 and month2 == month1 and date1 > date2:
        print("Check out date cannot be before check in date")
        check_out_date_wrong = True
    else:
        continue

check_out_date = f"{date2}/{month2}/{year2}"

while not_done:
    print("\nEnter Yes or No")
    booking = input("Do want to change the Booking Ref? ")
    if (booking.upper() == "YES"):
        booking_ref = input("Please enter the Booking Ref: ")
        not_done = False
    elif (booking.upper() == "NO"):
        booking_ref = "default booking_ref"
        not_done = False
    else:
        print("Please Enter a valid answer. Either yes or no")

while police_nr_not_done:
    print("\nEnter Yes or No")
    confirm = input("Do want to change the Policenr? ")
    if (confirm.upper() == "YES"):
        Policenr = input("Please enter the Booking Ref: ")
        police_nr_not_done = False
    elif (confirm.upper() == "NO"):
        Policenr = "default policenr number"
        police_nr_not_done = False
    else:
        print("Please Enter a valid answer. Either yes or no")

birthdate = input("Enter Birthdate: ")

ticket_to_Istanbul(name, check_in_date, check_out_date, booking_ref)
PrivatPolice(address, name, Policenr, birthdate)
agan_Hotel(name, date1, date2, abs(int(date1)-int(date2)), month1, month2)